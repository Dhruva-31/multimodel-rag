from docx import Document as DocxDocument
from models.document import Document
from models.metadata import Metadata
from ingestion.base_loader import BaseLoader


class DocxLoader(BaseLoader):

    def load(
        self,
        file_path: str,
        original_filename: str,
    ) -> list[Document]:

        doc = DocxDocument(file_path)

        documents = []

        paragraph_number = 1

        for p in doc.paragraphs:

            text = p.text.strip()

            if not text:
                continue

            documents.append(
                Document(
                    content=text,
                    metadata=Metadata(
                        source=file_path,
                        filename=original_filename,
                        type="docx",
                        paragraph=paragraph_number,
                    ),
                )
            )

            paragraph_number += 1

        return documents
